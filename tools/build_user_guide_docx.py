from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/Huong_dan_su_dung_PMEDIA_Dang_Bai_Tu_Dong_WordPress.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
LIGHT_YELLOW = "FFF7E6"


def main() -> int:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_header_footer(doc)
    add_title(doc)
    add_intro(doc)
    add_prepare_section(doc)
    add_connection_section(doc)
    add_excel_section(doc)
    add_image_section(doc)
    add_manual_publish_section(doc)
    add_schedule_section(doc)
    add_output_history_section(doc)
    add_troubleshooting_section(doc)
    add_safety_section(doc)
    add_quick_reference_section(doc)
    doc.save(OUTPUT)
    print(OUTPUT)
    return 0


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    set_style_font(normal, "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]
    for style_name, size, color, before, after in heading_tokens:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        set_style_font(style, "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        set_style_font(style, "Calibri")
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_style_font(style, name: str) -> None:
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_run_font(run, name: str = "Calibri", size: int | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.text = ""
    run = header.add_run("PMEDIA - ĐĂNG BÀI TỰ ĐỘNG WORDPRESS")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    run = footer.add_run("Tài liệu hướng dẫn sử dụng - kiểm tra preview trước khi đăng thật")
    set_run_font(run, size=9, color=MUTED)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Hướng dẫn sử dụng PMEDIA - ĐĂNG BÀI TỰ ĐỘNG WORDPRESS")
    set_run_font(run, size=24, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Quy trình chuẩn bị Excel, đặt tên hình, kết nối WordPress, đăng bài, hẹn lịch và kiểm tra kết quả"
    )
    set_run_font(run, size=12, color=MUTED)

    add_note(
        doc,
        "Mục tiêu",
        "Giúp người dùng đưa nội dung SEO HTML từ Excel lên WordPress đúng cấu trúc, có ảnh đại diện, ảnh trong bài, "
        "slug, mô tả meta, từ khóa Rank Math, file kết quả kèm link bài viết và lịch sử các lần chạy.",
        fill="F4F8FC",
    )


def add_intro(doc: Document) -> None:
    doc.add_heading("1. App dùng để làm gì?", level=1)
    add_para(
        doc,
        "PMEDIA - ĐĂNG BÀI TỰ ĐỘNG WORDPRESS là app desktop dùng để đọc file Excel, kiểm tra preview, upload ảnh local "
        "vào Media Library và tạo hoặc cập nhật bài viết WordPress qua REST API."
    )
    add_bullets(
        doc,
        [
            "Đọc sheet Bài SEO HTML trong Excel và giữ nguyên nội dung HTML.",
            "Đưa tiêu đề, slug, mô tả meta, danh mục, tags và từ khóa SEO vào WordPress.",
            "Upload ảnh đại diện/thumb và ảnh nội dung theo ma_bai.",
            "Chèn ảnh thumb làm hình đầu tiên trong bài, sau đó tới các ảnh nội dung _1, _2...",
            "Tự xuất bản sao Excel gốc kèm cột Link bài viết sau khi đăng xong.",
            "Lưu lịch sử các lần chạy trong tab Lịch sử.",
            "Hỗ trợ chạy thủ công hoặc hẹn lịch tự động.",
        ],
    )


def add_prepare_section(doc: Document) -> None:
    doc.add_heading("2. Chuẩn bị trước khi mở app", level=1)
    add_table(
        doc,
        ["Thành phần", "Đường dẫn hoặc yêu cầu"],
        [
            ["File chạy app", r"C:\Users\Admin\OneDrive\Tài liệu\Đăng bài tự động\dist\PMEDIA-DANG-BAI-TU-DONG-WORDPRESS.exe"],
            ["File Excel mẫu", r"docs\Mau_Excel_Dang_Bai_SEO_HTML_WordPress.xlsx"],
            ["Thư mục ảnh mẫu", r"docs\mau_ten_hinh_anh"],
            ["Quy ước tên hình", r"docs\Quy_uoc_dat_ten_hinh_anh.txt"],
            ["Tài khoản WordPress", "User có quyền tạo bài viết, upload media, tạo category/tag nếu cần."],
            ["Application Password", "Mật khẩu ứng dụng tạo trong WordPress, không phải mật khẩu đăng nhập thường."],
            ["Plugin Rank Math REST meta", r"wordpress\wordpress-auto-poster-rank-math-rest-meta.zip nếu cần ghi SEO fields Rank Math."],
        ],
        [2.1, 4.4],
    )
    add_note(
        doc,
        "Khuyến nghị",
        "Lần đầu chạy nên để status là draft, đăng thử 1-2 bài, kiểm tra trong WordPress rồi mới chạy toàn bộ file.",
        fill=LIGHT_YELLOW,
    )


def add_connection_section(doc: Document) -> None:
    doc.add_heading("3. Kết nối WordPress", level=1)
    add_para(
        doc,
        "App kết nối WordPress bằng REST API. Ba thông tin bắt buộc là URL website, username và Application Password."
    )
    add_table(
        doc,
        ["Ô trong app", "Cách nhập đúng", "Lỗi thường gặp"],
        [
            [
                "URL",
                "Nhập đầy đủ dạng https://tenmien.com. Không cần thêm /wp-json.",
                "Nếu chỉ nhập tenmien.com app sẽ báo No scheme supplied.",
            ],
            [
                "Username",
                "Nhập username WordPress, ví dụ admin. Không dùng tên hiển thị nếu khác username.",
                "Sai username có thể báo 401 hoặc rest_not_logged_in.",
            ],
            [
                "Application password",
                "Tạo trong WordPress Admin -> Users -> Profile -> Application Passwords.",
                "Không dùng mật khẩu đăng nhập thường. Sai password sẽ báo 401.",
            ],
        ],
        [1.35, 3.05, 2.1],
    )
    doc.add_heading("3.1. Cách tạo Application Password", level=2)
    add_numbers(
        doc,
        [
            "Vào trang quản trị WordPress, ví dụ https://tenmien.com/wp-admin.",
            "Đăng nhập bằng tài khoản có quyền tạo bài viết.",
            "Vào Users -> Profile hoặc Người dùng -> Hồ sơ.",
            "Kéo xuống phần Application Passwords hoặc Mật khẩu ứng dụng.",
            "Nhập tên ứng dụng, ví dụ PMEDIA Auto Poster.",
            "Bấm tạo mật khẩu mới và copy toàn bộ mật khẩu vừa sinh ra.",
            "Quay lại app, nhập URL, username, Application Password rồi bấm Kiểm tra kết nối.",
            "Chỉ đăng bài khi trạng thái kết nối báo OK.",
        ],
    )
    add_note(
        doc,
        "Lỗi 401/rest_not_logged_in",
        "Lỗi này nghĩa là WordPress chưa chấp nhận thông tin đăng nhập API. Hãy tạo lại Application Password đúng user, "
        "kiểm tra username, kiểm tra URL có https:// và kiểm tra plugin bảo mật/hosting có chặn REST API hay không.",
        fill=LIGHT_YELLOW,
    )
    doc.add_heading("3.2. Cài plugin metadata SEO cho Rank Math", level=2)
    add_para(
        doc,
        "Plugin này giúp WordPress REST API cho phép app ghi Tiêu đề SEO, Thẻ mô tả, Từ khóa chính và permalink vào "
        "Rank Math. Nên cài plugin trước khi chạy đăng bài chính thức, đặc biệt nếu sau khi test các ô Rank Math vẫn "
        "hiện trống trong phần Xem trước trình chỉnh sửa đoạn trích.",
    )
    add_numbers(
        doc,
        [
            "Mở WordPress Admin của website cần đăng bài.",
            "Vào Plugins -> Add New hoặc Gói mở rộng -> Cài mới.",
            "Bấm Upload Plugin hoặc Tải plugin lên.",
            "Chọn file wordpress/wordpress-auto-poster-rank-math-rest-meta.zip trong gói bàn giao.",
            "Bấm Install Now hoặc Cài đặt ngay.",
            "Sau khi cài xong, bấm Activate hoặc Kích hoạt.",
            "Quay lại app PMEDIA, bấm Kiểm tra kết nối rồi đăng thử 1 bài ở trạng thái draft.",
            "Mở bài draft trong WordPress và kiểm tra Rank Math đã có Tiêu đề SEO, Thẻ mô tả và Từ khóa chính.",
        ],
    )
    add_note(
        doc,
        "Khi nào bắt buộc cài plugin metadata?",
        "Nếu WordPress vẫn tạo được bài nhưng Rank Math không nhận SEO title, mô tả meta hoặc focus keyword, hãy cài "
        "plugin này rồi chạy lại. Với bài đã tồn tại, app sẽ cập nhật lại metadata SEO khi xử lý bài trùng.",
        fill=LIGHT_BLUE,
    )


def add_excel_section(doc: Document) -> None:
    doc.add_heading("4. Định dạng file Excel", level=1)
    add_para(
        doc,
        "App ưu tiên sheet có tên Bài SEO HTML. File mẫu đã có sẵn các cột khuyến nghị; người dùng chỉ cần sửa nội dung "
        "theo từng bài, giữ nguyên hàng tiêu đề và không đổi ý nghĩa cột."
    )
    add_table(
        doc,
        ["Cột Excel", "Bắt buộc", "App dùng để làm gì"],
        [
            ["ma_bai", "Nên có", "Nhận diện ảnh local thuộc bài nào."],
            ["Tiêu đề SEO", "Có", "Title WordPress và Rank Math title."],
            ["Nội dung HTML thuần", "Có", "Content WordPress. Giữ nguyên H2, H3, P, strong, liên hệ."],
            ["Slug", "Nên có", "Liên kết cố định/permalink của bài viết."],
            ["Mô tả Meta SEO", "Nên có", "Excerpt WordPress và Rank Math description."],
            ["Danh mục", "Tùy chọn", "Category WordPress. App tìm hoặc tạo nếu user có quyền."],
            ["tags", "Tùy chọn", "Thẻ WordPress, tách bằng dấu phẩy."],
            ["Từ khóa chính", "Nên có", "Rank Math focus keyword."],
            ["Từ khóa phụ đã phủ thêm", "Tùy chọn", "Gộp vào Rank Math focus keyword, tách bằng dấu phẩy."],
            ["status", "Nên có", "draft, publish, future hoặc private. Khuyến nghị dùng draft."],
            ["publish_date", "Tùy chọn", "Ngày đăng ISO, ví dụ 2026-07-10T09:00:00."],
            ["featured_image_url", "Tùy chọn", "URL ảnh đại diện từ internet nếu không dùng ảnh local."],
        ],
        [1.75, 0.85, 3.9],
    )
    add_note(
        doc,
        "Lưu ý về SEO fields",
        "App gửi rank_math_title, rank_math_description, rank_math_focus_keyword và rank_math_permalink vào WordPress. "
        "Nếu Rank Math chưa nhận các field này, hãy cài plugin hỗ trợ trong thư mục wordpress của dự án.",
    )
    doc.add_heading("4.1. Quy tắc nhập dữ liệu", level=2)
    add_bullets(
        doc,
        [
            "Không để trống Tiêu đề SEO và Nội dung HTML thuần.",
            "Slug nên viết không dấu, dùng dấu gạch ngang, ví dụ thong-cong-bien-hoa.",
            "Từ khóa phụ, tags nhập bằng dấu phẩy: từ khóa 1, từ khóa 2, từ khóa 3.",
            "Nếu chưa muốn public, giữ status là draft.",
            "Không chèn ảnh dung lượng quá lớn nếu không cần; app sẽ nén ảnh lớn khi upload.",
            "Sau khi đăng xong, app tự tạo file Excel kết quả với cột Link bài viết ở bên phải dữ liệu gốc.",
        ],
    )


def add_image_section(doc: Document) -> None:
    doc.add_heading("5. Định dạng và cách đặt tên hình", level=1)
    add_para(
        doc,
        "Ảnh local được nhận diện bằng ma_bai. Tên ảnh phải bắt đầu bằng đúng ma_bai trong Excel, sau đó là dấu gạch dưới."
    )
    add_table(
        doc,
        ["Mẫu tên hình", "Vai trò", "Ví dụ"],
        [
            ["{ma_bai}_bg.jpg", "Ảnh đại diện/thumb", "thong-cong-bien-hoa_bg.jpg"],
            ["{ma_bai}_thumb.webp", "Ảnh đại diện/thumb", "thong-cong-bien-hoa_thumb.webp"],
            ["{ma_bai}_thumbnail.png", "Ảnh đại diện/thumb", "thong-cong-bien-hoa_thumbnail.png"],
            ["{ma_bai}_featured.jpg", "Ảnh đại diện/thumb", "thong-cong-bien-hoa_featured.jpg"],
            ["{ma_bai}_1.jpg", "Ảnh nội dung thứ 1", "thong-cong-bien-hoa_1.jpg"],
            ["{ma_bai}_2.png", "Ảnh nội dung thứ 2", "thong-cong-bien-hoa_2.png"],
            ["{ma_bai}_3.webp", "Ảnh nội dung tiếp theo", "thong-cong-bien-hoa_3.webp"],
        ],
        [2.0, 2.1, 2.4],
    )
    add_bullets(
        doc,
        [
            "Đuôi file hỗ trợ: .jpg, .jpeg, .png, .webp.",
            "Ảnh _bg, _thumb, _thumbnail hoặc _featured được upload làm ảnh đại diện WordPress.",
            "Ảnh đại diện cũng được chèn làm hình đầu tiên trong nội dung bài.",
            "Ảnh _1, _2, _3 được chèn tiếp theo trong nội dung theo thứ tự số tăng dần.",
            "Số ảnh nội dung thực tế phụ thuộc ô Số ảnh tối đa mỗi bài trong app.",
            "File không bắt đầu bằng ma_bai hợp lệ sẽ hiện cảnh báo ảnh không khớp trong preview.",
        ],
    )
    add_note(
        doc,
        "Ví dụ đầy đủ",
        "Nếu ma_bai là thong-cong-bien-hoa thì bộ ảnh đúng là thong-cong-bien-hoa_bg.jpg, "
        "thong-cong-bien-hoa_1.jpg, thong-cong-bien-hoa_2.jpg. Không đặt tên kiểu ảnh 1.jpg hoặc thong cong bien hoa.jpg.",
        fill=LIGHT_YELLOW,
    )


def add_manual_publish_section(doc: Document) -> None:
    doc.add_heading("6. Đăng bài thủ công", level=1)
    add_numbers(
        doc,
        [
            "Mở file PMEDIA-DANG-BAI-TU-DONG-WORDPRESS.exe.",
            "Nhập URL, username và Application Password.",
            "Bấm Kiểm tra kết nối và chờ trạng thái OK.",
            "Chọn file Excel.",
            "Chọn thư mục ảnh nếu dùng ảnh local. Có thể bỏ trống nếu chỉ đăng nội dung HTML.",
            "Chỉnh Số ảnh tối đa chèn mỗi bài, Kích thước ảnh và Căn ảnh nếu cần.",
            "Kiểm tra bảng Preview: số bài, ma_bai, tiêu đề, category, tags, trạng thái, ảnh nền và số ảnh nội dung.",
            "Nếu có cảnh báo ảnh không khớp, kiểm tra lại tên file ảnh trước khi đăng.",
            "Bấm Đăng ngay và xác nhận.",
            "Sau khi hoàn tất, xem bảng Kết quả và tab Lịch sử.",
        ],
    )
    add_note(
        doc,
        "Bài trùng tiêu đề",
        "App đang bật cơ chế kiểm tra trùng tiêu đề. Nếu bài đã tồn tại, app ưu tiên cập nhật nội dung/ảnh/SEO của bài trùng "
        "thay vì tạo thêm một bài mới.",
    )


def add_schedule_section(doc: Document) -> None:
    doc.add_heading("7. Hẹn lịch tự động", level=1)
    add_para(
        doc,
        "Tab Lịch tự động cho phép chạy theo ngày, theo tuần, một lần theo ngày giờ hoặc custom cron. Lịch chạy theo thời gian thực của máy."
    )
    add_table(
        doc,
        ["Trường", "Cách dùng"],
        [
            ["File Excel", "Chọn file cố định dùng cho lịch."],
            ["Thư mục ảnh", "Chọn thư mục ảnh đặt tên theo ma_bai. Có thể bỏ trống."],
            ["Số ảnh tối đa mỗi bài", "Số ảnh nội dung _1, _2... được chèn sau ảnh thumb."],
            ["Tần suất", "Hàng ngày, hàng tuần, một lần theo ngày giờ hoặc custom cron."],
            ["Giờ chạy", "Dùng cho hàng ngày/hàng tuần, định dạng HH:mm."],
            ["Ngày giờ chạy", "Dùng cho lịch một lần."],
            ["Bật lịch tự động", "Tick để kích hoạt lịch."],
            ["Lưu lịch", "Lưu cấu hình lịch và cập nhật lần chạy kế tiếp."],
        ],
        [2.1, 4.4],
    )
    add_bullets(
        doc,
        [
            "Khi bấm nút X, app thu nhỏ xuống khay hệ thống và lịch vẫn chạy nền.",
            "Muốn thoát hẳn, bấm chuột phải vào icon khay hệ thống rồi chọn Thoát hẳn.",
            "Application Password không lưu vào settings. Khi mở app lại cần nhập lại password để lịch đăng được bài.",
            "Nếu dùng lịch một lần, thời gian chạy phải là thời gian tương lai.",
        ],
    )


def add_output_history_section(doc: Document) -> None:
    doc.add_heading("8. Kết quả, file xuất và lịch sử chạy", level=1)
    add_table(
        doc,
        ["Chức năng", "Kết quả tạo ra", "Vị trí xem"],
        [
            ["Xuất Excel gốc + link", "Bản sao file Excel gốc có thêm cột Link bài viết ở bên phải dữ liệu.", "Cùng thư mục với Excel gốc hoặc nơi người dùng chọn."],
            ["Bảng Kết quả", "Dòng, tiêu đề, trạng thái, link, ghi chú/lỗi.", "Tab Đăng thủ công."],
            ["Lịch sử chạy", "Thời gian, kiểu chạy, Excel, số thành công/lỗi/bỏ qua, file kết quả, chi tiết từng bài.", "Tab Lịch sử."],
            ["File lịch sử local", "Dữ liệu JSON dùng để app mở lại vẫn xem được lịch sử.", r"config\run_history.json"],
        ],
        [1.8, 3.25, 1.45],
    )
    add_bullets(
        doc,
        [
            "File Excel kết quả không sửa trực tiếp file gốc.",
            "Nếu bài đăng thành công, ô Link bài viết có hyperlink để mở bài phía người dùng.",
            "Nếu bài lỗi hoặc bị bỏ qua, dòng đó vẫn nằm trong lịch sử để kiểm tra nguyên nhân.",
            "config/run_history.json là dữ liệu local, không nên gửi công khai nếu có link website nội bộ.",
        ],
    )


def add_troubleshooting_section(doc: Document) -> None:
    doc.add_heading("9. Lỗi thường gặp và cách xử lý", level=1)
    add_table(
        doc,
        ["Hiện tượng", "Nguyên nhân thường gặp", "Cách xử lý"],
        [
            ["No scheme supplied", "URL thiếu https://", "Nhập lại URL dạng https://tenmien.com."],
            ["401/rest_not_logged_in", "Sai username hoặc Application Password", "Tạo lại Application Password đúng user, kiểm tra plugin bảo mật có chặn REST API không."],
            ["Không đọc được Excel", "Thiếu cột title/content hoặc Tiêu đề SEO/Nội dung HTML thuần", "Dùng lại file mẫu và giữ nguyên hàng tiêu đề."],
            ["Preview thiếu bài", "Có dòng trống title/content", "Kiểm tra các dòng trống hoặc ô nội dung bị thiếu."],
            ["Ảnh không lên WordPress", "Tên ảnh không khớp ma_bai hoặc sai thư mục", "Đặt lại tên theo {ma_bai}_bg, {ma_bai}_1, {ma_bai}_2."],
            ["SEO title/meta không hiện trong Rank Math", "Rank Math meta chưa show_in_rest", "Cài plugin wordpress-auto-poster-rank-math-rest-meta.zip."],
            ["Bài bị draft", "Cột status là draft", "Đây là khuyến nghị an toàn. Đổi thành publish nếu muốn public ngay."],
            ["Lịch không chạy", "Chưa tick bật lịch, app đã thoát hẳn hoặc chưa nhập password", "Bật lịch, để app ở khay hệ thống, nhập lại Application Password sau khi mở app."],
            ["Không thấy file kết quả", "Xuất Excel lỗi hoặc không có quyền ghi thư mục", "Xem log trong app, chọn nơi xuất khác bằng nút Xuất Excel gốc + link."],
        ],
        [1.7, 2.25, 2.55],
    )


def add_safety_section(doc: Document) -> None:
    doc.add_heading("10. Nguyên tắc an toàn khi đăng hàng loạt", level=1)
    add_bullets(
        doc,
        [
            "Luôn chạy thử 1-2 bài ở trạng thái draft trước khi đăng số lượng lớn.",
            "Kiểm tra preview trước khi bấm Đăng ngay.",
            "Không chia sẻ Application Password công khai.",
            "Không đóng app bằng Thoát hẳn khi đang đăng bài.",
            "Sau khi đăng xong, mở vài link trong file Excel kết quả để kiểm tra giao diện người dùng.",
            "Nếu có lỗi ảnh hoặc SEO, sửa Excel/tên ảnh rồi chạy lại. App sẽ cập nhật bài trùng tiêu đề.",
        ],
    )


def add_quick_reference_section(doc: Document) -> None:
    doc.add_heading("11. Tóm tắt thao tác nhanh", level=1)
    add_table(
        doc,
        ["Mục tiêu", "Thao tác nhanh"],
        [
            ["Đăng bài có ảnh", "Chọn Excel -> chọn thư mục ảnh -> Preview -> Đăng ngay -> kiểm tra Link bài viết."],
            ["Chỉ kiểm tra dữ liệu", "Chọn Excel -> Preview -> sửa cảnh báo nếu có -> chưa bấm Đăng ngay."],
            ["Hẹn giờ đăng", "Tab Lịch tự động -> chọn Excel/ảnh -> chọn giờ -> bật lịch -> lưu lịch."],
            ["Xem lần chạy cũ", "Mở tab Lịch sử -> chọn dòng -> xem chi tiết bên dưới."],
            ["Tạo lại file mẫu", "Chạy tools/create_excel_image_template.py để xuất Excel mẫu và ảnh mẫu."],
        ],
        [2.0, 4.5],
    )
    add_note(
        doc,
        "Phiên bản tài liệu",
        f"Cập nhật ngày {datetime.now().strftime('%d/%m/%Y')} theo bản app có xuất Excel kèm link bài viết và lưu lịch sử chạy.",
        fill=LIGHT_BLUE,
    )


def add_para(doc: Document, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(item))


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        set_run_font(paragraph.add_run(item))


def add_note(doc: Document, label: str, body: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, bold=True, color=DARK_BLUE)
    run = paragraph.add_run(body)
    set_run_font(run)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_shading(header_cells[index], LIGHT_BLUE)
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, bold=True, color=DARK_BLUE)

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_width(cells[index], widths[index])
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=10 if len(value) > 90 else 11)
    doc.add_paragraph()


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            set_cell_width(row.cells[index], width)
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[index])

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    width = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


if __name__ == "__main__":
    raise SystemExit(main())
